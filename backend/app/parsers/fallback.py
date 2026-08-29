from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document as DocxDocument

from app.core.models import BlockKind, DocumentBlock, ParsedDocument
from .base import DocumentParser


class FallbackParser(DocumentParser):
    def parse(self, source_path: Path, work_dir: Path) -> ParsedDocument:
        suffix = source_path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return self._parse_text(source_path)
        if suffix == ".pdf":
            return self._parse_pdf(source_path)
        if suffix == ".docx":
            return self._parse_docx(source_path)
        raise RuntimeError(
            f"No fallback parser for {suffix}. Install MinerU for scanned images and rich office documents."
        )

    def _parse_text(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8")
        chunks = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        blocks = [
            DocumentBlock(id=f"b{i:05d}", kind=BlockKind.TEXT, text=chunk, page=0)
            for i, chunk in enumerate(chunks)
        ]
        return ParsedDocument(str(path), blocks, parser="fallback-text")

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        doc = fitz.open(path)
        blocks: list[DocumentBlock] = []
        idx = 0
        for page_index, page in enumerate(doc):
            page_dict = page.get_text("dict")
            for item in page_dict.get("blocks", []):
                if "lines" not in item:
                    continue
                parts: list[str] = []
                for line in item.get("lines", []):
                    for span in line.get("spans", []):
                        parts.append(span.get("text", ""))
                text = " ".join(x.strip() for x in parts if x.strip()).strip()
                if not text:
                    continue
                rect = page.rect
                bbox = item.get("bbox")
                normalized = None
                if bbox and rect.width and rect.height:
                    normalized = [
                        bbox[0] / rect.width * 1000,
                        bbox[1] / rect.height * 1000,
                        bbox[2] / rect.width * 1000,
                        bbox[3] / rect.height * 1000,
                    ]
                blocks.append(
                    DocumentBlock(
                        id=f"b{idx:05d}", kind=BlockKind.TEXT, text=text,
                        page=page_index, bbox=normalized,
                    )
                )
                idx += 1
        doc.close()
        return ParsedDocument(
            str(path), blocks, parser="fallback-pymupdf",
            warnings=["MinerU was unavailable; formulas/images/layout may be incomplete in fallback PDF parsing."],
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(path)
        blocks: list[DocumentBlock] = []
        idx = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            kind = BlockKind.TITLE if paragraph.style and paragraph.style.name.startswith("Heading") else BlockKind.TEXT
            blocks.append(DocumentBlock(id=f"b{idx:05d}", kind=kind, text=text, page=0))
            idx += 1
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            blocks.append(DocumentBlock(id=f"b{idx:05d}", kind=BlockKind.TABLE, text="\n".join(rows), page=0))
            idx += 1
        return ParsedDocument(
            str(path), blocks, parser="fallback-python-docx",
            warnings=["MinerU was unavailable; embedded drawings/equations may not be reconstructed by the DOCX fallback."],
        )
